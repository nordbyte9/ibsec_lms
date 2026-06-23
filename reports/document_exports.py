from __future__ import annotations

from dataclasses import dataclass, replace
from datetime import date, datetime
from decimal import Decimal
from io import BytesIO
from math import isfinite
from typing import Mapping, Sequence
from xml.etree import ElementTree
from zipfile import BadZipFile, ZipFile

from django.utils import timezone
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

XLSX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
)
DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)

EXCEL_CELL_TEXT_LIMIT = 32_767


@dataclass(frozen=True)
class CertificateData:
    certificate_number: str
    organization_name: str
    employee_name: str
    username: str
    department: str
    position: str
    course_title: str
    program_title: str
    completed_at: date
    result_percent: float | None
    valid_until: date | None
    signer_title: str
    signer_name: str


def _is_valid_xml_character(character: str) -> bool:
    """Return whether a character is permitted by XML 1.0."""

    codepoint = ord(character)
    if character in {"\t", "\n", "\r"}:
        return True
    if codepoint < 0x20:
        return False
    if 0xD800 <= codepoint <= 0xDFFF:
        return False
    if (codepoint & 0xFFFF) in {0xFFFE, 0xFFFF}:
        return False
    return codepoint <= 0x10FFFF


def _clean_xml_text(value: object) -> str:
    if value is None:
        return ""
    return "".join(
        character
        for character in str(value)
        if _is_valid_xml_character(character)
    )


def _safe_docx_text(value: object) -> str:
    return _clean_xml_text(value)


def _safe_excel_text(value: object) -> object:
    """Return XML-safe text and prevent spreadsheet formula injection."""

    if value is None:
        return ""
    if not isinstance(value, str):
        return value

    cleaned = _clean_xml_text(value)[:EXCEL_CELL_TEXT_LIMIT]
    candidate = cleaned.lstrip(" \t\r\n")
    if cleaned.startswith(("\t", "\r", "\n")) or candidate.startswith(
        ("=", "+", "-", "@")
    ):
        return f"'{cleaned}"
    return cleaned


def _excel_datetime(value: datetime | None) -> datetime | None:
    if value is None:
        return None
    if timezone.is_aware(value):
        value = timezone.localtime(value)
    return value.replace(tzinfo=None)


def _safe_excel_number(value: object) -> int | float | None:
    if value is None or isinstance(value, bool):
        return None
    if isinstance(value, Decimal):
        value = float(value)
    if isinstance(value, int):
        return value
    if isinstance(value, float):
        return value if isfinite(value) else None
    return None


def _apply_cell_border(cell, color: str = "D9E2F3") -> None:
    thin = Side(style="thin", color=color)
    cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)


def _validate_zip_xml_package(
    content: bytes,
    *,
    required_entries: Sequence[str],
) -> None:
    """Validate ZIP CRC and XML syntax for an Office Open XML package."""

    if not content.startswith(b"PK"):
        raise ValueError("Сформированный документ не является ZIP-пакетом OOXML")

    try:
        with ZipFile(BytesIO(content)) as package:
            damaged_entry = package.testzip()
            if damaged_entry:
                raise ValueError(
                    f"Повреждён элемент OOXML-пакета: {damaged_entry}"
                )

            names = set(package.namelist())
            missing = [name for name in required_entries if name not in names]
            if missing:
                raise ValueError(
                    "В OOXML-пакете отсутствуют обязательные элементы: "
                    + ", ".join(missing)
                )

            for name in names:
                if name.endswith((".xml", ".rels")):
                    try:
                        ElementTree.fromstring(package.read(name))
                    except ElementTree.ParseError as exc:
                        raise ValueError(
                            f"Некорректный XML в элементе {name}"
                        ) from exc
    except BadZipFile as exc:
        raise ValueError("Сформированный OOXML-пакет повреждён") from exc


def _validate_xlsx(content: bytes) -> None:
    _validate_zip_xml_package(
        content,
        required_entries=(
            "[Content_Types].xml",
            "xl/workbook.xml",
            "xl/worksheets/sheet1.xml",
            "xl/worksheets/sheet2.xml",
        ),
    )

    workbook = load_workbook(BytesIO(content), read_only=True, data_only=False)
    try:
        expected_sheets = ["Реестр обучения", "Сводка"]
        if workbook.sheetnames != expected_sheets:
            raise ValueError("Сформирован XLSX с неожиданным набором листов")
    finally:
        workbook.close()


def _validate_docx(content: bytes) -> None:
    _validate_zip_xml_package(
        content,
        required_entries=(
            "[Content_Types].xml",
            "word/document.xml",
            "word/styles.xml",
        ),
    )
    Document(BytesIO(content))


def build_training_registry_xlsx(
    rows: Sequence[Mapping[str, object]],
    *,
    generated_at: datetime,
    filters_description: str = "Все данные",
) -> bytes:
    """Build an Excel-compatible XLSX training registry.

    The workbook intentionally uses only standard worksheet cells, styles and
    one worksheet autofilter. It does not create an overlapping Excel table or
    DrawingML chart, which improves compatibility with desktop Excel versions.
    """

    normalized_rows = list(rows)

    workbook = Workbook()
    workbook.properties.creator = "IBSec LMS"
    workbook.properties.title = (
        "Реестр обучения по информационной безопасности"
    )
    workbook.properties.subject = "Отчёт о назначениях и результатах обучения"
    workbook.properties.keywords = "IBSec LMS, обучение, отчёт, реестр"

    registry = workbook.active
    registry.title = "Реестр обучения"

    title_fill = PatternFill("solid", fgColor="1F4E78")
    header_fill = PatternFill("solid", fgColor="D9EAF7")
    accent_fill = PatternFill("solid", fgColor="E2F0D9")
    overdue_fill = PatternFill("solid", fgColor="FCE4D6")
    alternate_fill = PatternFill("solid", fgColor="F7FAFC")
    white_font = Font(color="FFFFFF", bold=True, size=14)
    header_font = Font(bold=True, color="1F1F1F")
    centered = Alignment(
        horizontal="center",
        vertical="center",
        wrap_text=True,
    )
    wrapped = Alignment(vertical="top", wrap_text=True)

    headers = [
        "ID назначения",
        "Сотрудник",
        "Логин",
        "Подразделение",
        "Должность",
        "Курс",
        "Программа ИБ",
        "Статус",
        "Дата назначения",
        "Срок прохождения",
        "Дата завершения",
        "Последний результат, %",
        "Тест пройден",
        "Количество попыток",
        "Действительно до",
    ]

    registry.merge_cells(
        start_row=1,
        start_column=1,
        end_row=1,
        end_column=len(headers),
    )
    title_cell = registry.cell(
        1,
        1,
        "Реестр обучения по информационной безопасности",
    )
    title_cell.fill = title_fill
    title_cell.font = white_font
    title_cell.alignment = centered
    registry.row_dimensions[1].height = 28

    generated_local = _excel_datetime(generated_at)
    registry.cell(2, 1, "Сформирован")
    registry.cell(2, 1).font = header_font
    registry.cell(2, 2, generated_local)
    registry.cell(2, 2).number_format = "dd.mm.yyyy hh:mm"

    registry.cell(3, 1, "Фильтры")
    registry.cell(3, 1).font = header_font
    registry.merge_cells(
        start_row=3,
        start_column=2,
        end_row=3,
        end_column=len(headers),
    )
    registry.cell(3, 2, _safe_excel_text(filters_description))
    registry.cell(3, 2).alignment = wrapped

    header_row = 5
    for col_index, header in enumerate(headers, start=1):
        cell = registry.cell(header_row, col_index, header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = centered
        _apply_cell_border(cell)
    registry.row_dimensions[header_row].height = 42

    status_counts: dict[str, int] = {}
    percent_values: list[float] = []

    for row_index, item in enumerate(
        normalized_rows,
        start=header_row + 1,
    ):
        latest_percent = _safe_excel_number(item.get("latest_percent"))
        attempts = _safe_excel_number(item.get("attempts"))
        assignment_id = _safe_excel_number(item.get("assignment_id"))

        values = [
            assignment_id,
            _safe_excel_text(item.get("employee", "")),
            _safe_excel_text(item.get("username", "")),
            _safe_excel_text(item.get("department", "")),
            _safe_excel_text(item.get("position", "")),
            _safe_excel_text(item.get("course", "")),
            _safe_excel_text(item.get("program", "")),
            _safe_excel_text(item.get("status", "")),
            _excel_datetime(item.get("assigned_at")),
            item.get("due_date"),
            _excel_datetime(item.get("completed_at")),
            latest_percent,
            _safe_excel_text(item.get("passed", "")),
            attempts if attempts is not None else 0,
            item.get("valid_until"),
        ]

        for col_index, value in enumerate(values, start=1):
            cell = registry.cell(row_index, col_index, value)
            cell.alignment = (
                centered
                if col_index in {1, 8, 9, 10, 11, 12, 13, 14, 15}
                else wrapped
            )
            if row_index % 2 == 0:
                cell.fill = alternate_fill
            _apply_cell_border(cell, "E7E6E6")

        for column in (9, 11):
            registry.cell(row_index, column).number_format = (
                "dd.mm.yyyy hh:mm"
            )
        for column in (10, 15):
            registry.cell(row_index, column).number_format = "dd.mm.yyyy"
        registry.cell(row_index, 12).number_format = "0.00"

        status = _clean_xml_text(item.get("status", ""))
        status_counts[status] = status_counts.get(status, 0) + 1
        if status == "Завершено":
            registry.cell(row_index, 8).fill = accent_fill
        elif status == "Просрочено":
            registry.cell(row_index, 8).fill = overdue_fill

        if isinstance(latest_percent, (int, float)):
            percent_values.append(float(latest_percent))

    filter_last_row = (
        header_row + len(normalized_rows)
        if normalized_rows
        else header_row
    )
    registry.freeze_panes = "A6"
    registry.auto_filter.ref = f"A{header_row}:O{filter_last_row}"
    registry.sheet_view.showGridLines = False
    registry.page_setup.orientation = "landscape"
    registry.page_setup.fitToWidth = 1
    registry.page_setup.fitToHeight = 0
    registry.sheet_properties.pageSetUpPr.fitToPage = True

    if not normalized_rows:
        registry.merge_cells(
            start_row=header_row + 1,
            start_column=1,
            end_row=header_row + 1,
            end_column=len(headers),
        )
        empty_cell = registry.cell(
            header_row + 1,
            1,
            "По выбранным фильтрам данные отсутствуют",
        )
        empty_cell.alignment = centered
        empty_cell.fill = alternate_fill
        _apply_cell_border(empty_cell, "E7E6E6")

    widths = {
        1: 14,
        2: 26,
        3: 18,
        4: 24,
        5: 24,
        6: 34,
        7: 30,
        8: 16,
        9: 19,
        10: 18,
        11: 19,
        12: 21,
        13: 15,
        14: 19,
        15: 18,
    }
    for column_index, width in widths.items():
        registry.column_dimensions[get_column_letter(column_index)].width = width

    summary = workbook.create_sheet("Сводка")
    summary.sheet_view.showGridLines = False
    summary.merge_cells("A1:F1")
    summary["A1"] = "Сводка по реестру обучения"
    summary["A1"].fill = title_fill
    summary["A1"].font = white_font
    summary["A1"].alignment = centered
    summary.row_dimensions[1].height = 28

    completed = status_counts.get("Завершено", 0)
    overdue = status_counts.get("Просрочено", 0)
    in_progress = status_counts.get("В процессе", 0)
    assigned = status_counts.get("Назначено", 0)
    average_percent = (
        round(sum(percent_values) / len(percent_values), 2)
        if percent_values
        else 0
    )

    metrics = [
        ("Всего назначений", len(normalized_rows)),
        ("Завершено", completed),
        ("Просрочено", overdue),
        ("Средний результат, %", average_percent),
    ]
    for row_index, (label, value) in enumerate(metrics, start=3):
        summary.cell(row_index, 1, label)
        summary.cell(row_index, 2, value)
        summary.cell(row_index, 1).font = header_font
        summary.cell(row_index, 1).fill = header_fill
        summary.cell(row_index, 2).fill = accent_fill
        summary.cell(row_index, 1).alignment = wrapped
        summary.cell(row_index, 2).alignment = centered
        _apply_cell_border(summary.cell(row_index, 1))
        _apply_cell_border(summary.cell(row_index, 2))
    summary["B6"].number_format = "0.00"

    summary["A9"] = "Статус"
    summary["B9"] = "Количество"
    for cell in summary[9][:2]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = centered
        _apply_cell_border(cell)

    status_rows = [
        ("Назначено", assigned),
        ("В процессе", in_progress),
        ("Завершено", completed),
        ("Просрочено", overdue),
    ]
    for row_index, (label, value) in enumerate(status_rows, start=10):
        summary.cell(row_index, 1, label)
        summary.cell(row_index, 2, value)
        summary.cell(row_index, 1).alignment = wrapped
        summary.cell(row_index, 2).alignment = centered
        _apply_cell_border(summary.cell(row_index, 1))
        _apply_cell_border(summary.cell(row_index, 2))

    summary["D3"] = "Примечание"
    summary["D3"].font = header_font
    summary["D3"].fill = header_fill
    summary["D3"].alignment = centered
    _apply_cell_border(summary["D3"])
    summary.merge_cells("D4:F8")
    summary["D4"] = (
        "Сводка сформирована без встроенной диаграммы для максимальной "
        "совместимости с настольными версиями Microsoft Excel."
    )
    summary["D4"].alignment = Alignment(
        vertical="top",
        wrap_text=True,
    )
    _apply_cell_border(summary["D4"])

    summary.column_dimensions["A"].width = 28
    summary.column_dimensions["B"].width = 18
    summary.column_dimensions["C"].width = 3
    summary.column_dimensions["D"].width = 22
    summary.column_dimensions["E"].width = 16
    summary.column_dimensions["F"].width = 16
    summary.page_setup.orientation = "landscape"
    summary.page_setup.fitToWidth = 1
    summary.page_setup.fitToHeight = 1
    summary.sheet_properties.pageSetUpPr.fitToPage = True

    workbook.active = 0

    output = BytesIO()
    workbook.save(output)
    content = output.getvalue()
    _validate_xlsx(content)
    return content


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    table_header = OxmlElement("w:tblHeader")
    table_header.set(qn("w:val"), "true")
    tr_pr.append(table_header)


def _set_run_font(run, *, size: int | None = None) -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    if size is not None:
        run.font.size = Pt(size)


def build_certificate_docx(data: CertificateData) -> bytes:
    """Build and validate a DOCX certificate."""

    data = replace(
        data,
        certificate_number=_safe_docx_text(data.certificate_number),
        organization_name=_safe_docx_text(data.organization_name),
        employee_name=_safe_docx_text(data.employee_name),
        username=_safe_docx_text(data.username),
        department=_safe_docx_text(data.department),
        position=_safe_docx_text(data.position),
        course_title=_safe_docx_text(data.course_title),
        program_title=_safe_docx_text(data.program_title),
        signer_title=_safe_docx_text(data.signer_title),
        signer_name=_safe_docx_text(data.signer_name),
    )

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal_style.font.size = Pt(11)

    organization = document.add_paragraph()
    organization.alignment = WD_ALIGN_PARAGRAPH.CENTER
    organization_run = organization.add_run(data.organization_name)
    organization_run.bold = True
    _set_run_font(organization_run, size=13)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(10)
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run("СЕРТИФИКАТ")
    title_run.bold = True
    _set_run_font(title_run, size=26)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(
        "о прохождении обучения по информационной безопасности"
    )
    subtitle_run.italic = True
    _set_run_font(subtitle_run, size=12)

    number = document.add_paragraph()
    number.alignment = WD_ALIGN_PARAGRAPH.CENTER
    number_run = number.add_run(f"№ {data.certificate_number}")
    number_run.bold = True
    _set_run_font(number_run)

    intro = document.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro.paragraph_format.space_before = Pt(14)
    intro.paragraph_format.space_after = Pt(6)
    intro_run = intro.add_run("Настоящим подтверждается, что")
    _set_run_font(intro_run, size=12)

    employee = document.add_paragraph()
    employee.alignment = WD_ALIGN_PARAGRAPH.CENTER
    employee.paragraph_format.space_after = Pt(10)
    employee_run = employee.add_run(data.employee_name)
    employee_run.bold = True
    _set_run_font(employee_run, size=18)

    course_text = document.add_paragraph()
    course_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prefix_run = course_text.add_run("успешно завершил(а) курс\n")
    _set_run_font(prefix_run)
    course_run = course_text.add_run(f"«{data.course_title}»")
    course_run.bold = True
    _set_run_font(course_run, size=15)

    details = document.add_table(rows=1, cols=2)
    details.style = "Table Grid"
    details.autofit = False
    details.columns[0].width = Cm(6)
    details.columns[1].width = Cm(10)
    _set_repeat_table_header(details.rows[0])
    details.rows[0].cells[0].text = "Показатель"
    details.rows[0].cells[1].text = "Значение"

    for cell in details.rows[0].cells:
        _set_cell_shading(cell, "D9EAF7")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                _set_run_font(run)

    detail_rows = [
        ("Учетная запись", data.username),
        ("Подразделение", data.department or "—"),
        ("Должность", data.position or "—"),
        ("Программа обучения", data.program_title or "—"),
        ("Дата завершения", data.completed_at.strftime("%d.%m.%Y")),
        (
            "Результат тестирования",
            (
                f"{data.result_percent:.2f}%"
                if data.result_percent is not None
                else "Не указан"
            ),
        ),
        (
            "Срок действия",
            (
                f"до {data.valid_until.strftime('%d.%m.%Y')}"
                if data.valid_until
                else "Не ограничен"
            ),
        ),
    ]

    for label, value in detail_rows:
        cells = details.add_row().cells
        cells[0].text = _safe_docx_text(label)
        cells[1].text = _safe_docx_text(value)
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    _set_run_font(run)
        if cells[0].paragraphs and cells[0].paragraphs[0].runs:
            cells[0].paragraphs[0].runs[0].bold = True

    document.add_paragraph()

    signatures = document.add_table(rows=2, cols=2)
    signatures.autofit = False
    signatures.columns[0].width = Cm(9)
    signatures.columns[1].width = Cm(7)
    signatures.cell(0, 0).text = (
        data.signer_title
        or "Ответственный за информационную безопасность"
    )
    signatures.cell(0, 1).text = data.signer_name or "________________________"
    signatures.cell(1, 0).text = "Подпись"
    signatures.cell(1, 1).text = "Ф.И.О."

    for row in signatures.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    _set_run_font(run, size=10)

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(12)
    note_run = note.add_run(
        "Документ сформирован информационной системой IBSec LMS.\n"
        "Достоверность сведений подтверждается данными журнала обучения."
    )
    note_run.italic = True
    _set_run_font(note_run, size=8)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(
        f"IBSec LMS • Сертификат {data.certificate_number}"
    )
    _set_run_font(footer_run, size=8)

    properties = document.core_properties
    properties.title = f"Сертификат {data.certificate_number}"
    properties.subject = (
        "Подтверждение прохождения обучения по информационной безопасности"
    )
    properties.author = "IBSec LMS"
    properties.keywords = (
        "IBSec LMS, обучение, информационная безопасность, сертификат"
    )

    output = BytesIO()
    document.save(output)
    content = output.getvalue()
    _validate_docx(content)
    return content
