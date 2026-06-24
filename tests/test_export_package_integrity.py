from datetime import date, datetime, timezone
from io import BytesIO
from xml.etree import ElementTree
from zipfile import ZipFile

from django.test import SimpleTestCase
from docx import Document
from openpyxl import load_workbook

from reports.document_exports import (
    CertificateData,
    build_certificate_docx,
    build_training_registry_xlsx,
)


class ExportPackageIntegrityTests(SimpleTestCase):
    def _assert_valid_ooxml_zip(self, content: bytes) -> set[str]:
        self.assertTrue(content.startswith(b"PK"))

        with ZipFile(BytesIO(content)) as package:
            self.assertIsNone(package.testzip())
            names = set(package.namelist())

            for name in names:
                if name.endswith((".xml", ".rels")):
                    ElementTree.fromstring(package.read(name))

        return names

    def test_xlsx_is_simple_excel_compatible_package(self):
        content = build_training_registry_xlsx(
            [
                {
                    "assignment_id": 1,
                    "employee": "Иван Петров",
                    "username": "employee",
                    "department": "ИТ",
                    "position": "Аналитик",
                    "course": "Основы ИБ",
                    "program": "Базовая программа",
                    "status": "Завершено",
                    "assigned_at": datetime(
                        2026,
                        6,
                        1,
                        10,
                        30,
                        tzinfo=timezone.utc,
                    ),
                    "due_date": date(2026, 6, 30),
                    "completed_at": datetime(
                        2026,
                        6,
                        20,
                        12,
                        0,
                        tzinfo=timezone.utc,
                    ),
                    "latest_percent": 90.0,
                    "passed": "Да",
                    "attempts": 1,
                    "valid_until": date(2027, 6, 20),
                }
            ],
            generated_at=datetime(
                2026,
                6,
                23,
                9,
                0,
                tzinfo=timezone.utc,
            ),
        )

        names = self._assert_valid_ooxml_zip(content)

        # Compatibility mode intentionally has no DrawingML chart or Excel
        # table parts. The formatted cells and worksheet autofilter remain.
        self.assertFalse(any(name.startswith("xl/drawings/") for name in names))
        self.assertFalse(any(name.startswith("xl/tables/") for name in names))

        workbook = load_workbook(BytesIO(content), data_only=False)
        try:
            self.assertEqual(
                workbook.sheetnames,
                ["Реестр обучения", "Сводка"],
            )
            registry = workbook["Реестр обучения"]
            self.assertEqual(registry["A5"].value, "ID назначения")
            self.assertEqual(registry["B6"].value, "Иван Петров")
            self.assertEqual(registry.auto_filter.ref, "A5:O6")
        finally:
            workbook.close()

    def test_xlsx_sanitizes_formula_and_invalid_xml_characters(self):
        content = build_training_registry_xlsx(
            [
                {
                    "assignment_id": 1,
                    "employee": '=HYPERLINK("https://invalid.example")',
                    "username": "employee",
                    "department": "ИТ\x01",
                    "position": "Аналитик",
                    "course": "Основы ИБ",
                    "program": "",
                    "status": "Назначено",
                    "assigned_at": None,
                    "due_date": None,
                    "completed_at": None,
                    "latest_percent": None,
                    "passed": "—",
                    "attempts": 0,
                    "valid_until": None,
                }
            ],
            generated_at=datetime(
                2026,
                6,
                23,
                9,
                0,
                tzinfo=timezone.utc,
            ),
            filters_description="Все данные\x00",
        )

        self._assert_valid_ooxml_zip(content)
        workbook = load_workbook(BytesIO(content), data_only=False)
        try:
            registry = workbook["Реестр обучения"]
            self.assertEqual(
                registry["B6"].value,
                "'=HYPERLINK(\"https://invalid.example\")",
            )
            self.assertEqual(registry["D6"].value, "ИТ")
            self.assertEqual(registry["B3"].value, "Все данные")
        finally:
            workbook.close()

    def test_empty_xlsx_is_valid(self):
        content = build_training_registry_xlsx(
            [],
            generated_at=datetime(
                2026,
                6,
                23,
                9,
                0,
                tzinfo=timezone.utc,
            ),
        )

        self._assert_valid_ooxml_zip(content)
        workbook = load_workbook(BytesIO(content), data_only=False)
        try:
            self.assertIn(
                "По выбранным фильтрам данные отсутствуют",
                workbook["Реестр обучения"]["A6"].value,
            )
        finally:
            workbook.close()

    def test_docx_is_valid_and_sanitized(self):
        content = build_certificate_docx(
            CertificateData(
                certificate_number="IBSEC-000001",
                organization_name="IBSec LMS",
                employee_name="Иван\x01 Петров",
                username="employee",
                department="ИТ",
                position="Аналитик",
                course_title="Основы информационной безопасности",
                program_title="Базовая программа ИБ",
                completed_at=date(2026, 6, 20),
                result_percent=90.0,
                valid_until=date(2027, 6, 20),
                signer_title="Ответственный за ИБ",
                signer_name="Ольга Соколова",
            )
        )

        self._assert_valid_ooxml_zip(content)
        document = Document(BytesIO(content))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        table_text = "\n".join(
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        )

        self.assertIn("Иван Петров", text)
        self.assertNotIn("\x01", text)
        self.assertIn("90.00%", table_text)
