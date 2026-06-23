from datetime import timedelta
from io import BytesIO

from django.test import TestCase
from django.urls import reverse
from django.utils import timezone
from docx import Document
from openpyxl import load_workbook

from assignments.models import CourseAssignment
from audit.models import AuditLog
from courses.models import Course, SecurityCategory, TrainingProgram
from quizzes.models import Quiz, Submission
from tests.utils import create_user


class DocumentExportTests(TestCase):
    password = 'password123'

    def setUp(self):
        self.officer = create_user(
            'officer',
            password=self.password,
            role='security_officer',
            first_name='Ольга',
            last_name='Соколова',
        )
        self.employee = create_user(
            'employee',
            password=self.password,
            role='employee',
            department_name='ИТ',
            position_name='Аналитик',
            first_name='Иван',
            last_name='Петров',
        )
        self.other_employee = create_user(
            'other',
            password=self.password,
            role='employee',
            first_name='Анна',
            last_name='Смирнова',
        )
        category = SecurityCategory.objects.create(code='basic', name='Основы ИБ')
        program = TrainingProgram.objects.create(
            title='Базовая программа ИБ',
            category=category,
            is_mandatory=True,
        )
        self.course = Course.objects.create(
            title='Основы информационной безопасности',
            training_program=program,
            author=self.officer,
            is_published=True,
            validity_days=365,
        )
        self.assignment = CourseAssignment.objects.create(
            employee=self.employee,
            course=self.course,
            assigned_by=self.officer,
            due_date=timezone.localdate() + timedelta(days=10),
            status=CourseAssignment.Status.COMPLETED,
        )
        self.active_assignment = CourseAssignment.objects.create(
            employee=self.other_employee,
            course=self.course,
            assigned_by=self.officer,
            due_date=timezone.localdate() + timedelta(days=20),
            status=CourseAssignment.Status.IN_PROGRESS,
        )
        quiz = Quiz.objects.create(course=self.course, title='Итоговый тест')
        Submission.objects.create(
            user=self.employee,
            quiz=quiz,
            score=9,
            percent=90.0,
            passed=True,
            attempt_number=1,
        )

    def test_registry_requires_authentication(self):
        response = self.client.get(reverse('reports:training_registry_xlsx'))
        self.assertEqual(response.status_code, 302)

    def test_employee_cannot_export_registry(self):
        self.client.login(username=self.employee.username, password=self.password)
        response = self.client.get(reverse('reports:training_registry_xlsx'))
        self.assertEqual(response.status_code, 403)

    def test_officer_exports_valid_registry_and_audit_entry(self):
        self.client.login(username=self.officer.username, password=self.password)
        response = self.client.get(reverse('reports:training_registry_xlsx'))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response['Content-Type'],
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        )
        workbook = load_workbook(BytesIO(response.content), data_only=False)
        self.assertEqual(workbook.sheetnames, ['Реестр обучения', 'Сводка'])
        registry = workbook['Реестр обучения']
        self.assertEqual(registry['A5'].value, 'ID назначения')
        self.assertEqual(registry['B6'].value, 'Иван Петров')
        self.assertEqual(registry['F6'].value, self.course.title)
        self.assertTrue(
            AuditLog.objects.filter(user=self.officer, action='xlsx_export').exists()
        )

    def test_employee_downloads_own_completed_certificate(self):
        self.client.login(username=self.employee.username, password=self.password)
        response = self.client.get(
            reverse('reports:assignment_certificate_docx', args=[self.assignment.pk])
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response['Content-Type'],
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        document = Document(BytesIO(response.content))
        document_text = '\n'.join(paragraph.text for paragraph in document.paragraphs)
        table_text = '\n'.join(
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        )
        self.assertIn('Иван Петров', document_text)
        self.assertIn(self.course.title, document_text)
        self.assertIn('90.00%', table_text)
        self.assertTrue(
            AuditLog.objects.filter(user=self.employee, action='docx_certificate').exists()
        )

    def test_employee_cannot_download_another_users_certificate(self):
        self.client.login(username=self.other_employee.username, password=self.password)
        response = self.client.get(
            reverse('reports:assignment_certificate_docx', args=[self.assignment.pk])
        )
        self.assertEqual(response.status_code, 403)

    def test_incomplete_assignment_has_no_certificate(self):
        self.client.login(username=self.other_employee.username, password=self.password)
        response = self.client.get(
            reverse('reports:assignment_certificate_docx', args=[self.active_assignment.pk])
        )
        self.assertEqual(response.status_code, 404)

    def test_django_superuser_can_export_registry(self):
        superuser = create_user(
            'root',
            password=self.password,
            role='employee',
        )
        superuser.is_superuser = True
        superuser.is_staff = True
        superuser.save(update_fields=['is_superuser', 'is_staff'])
        self.client.login(username=superuser.username, password=self.password)

        response = self.client.get(reverse('reports:training_registry_xlsx'))

        self.assertEqual(response.status_code, 200)
        self.assertTrue(
            AuditLog.objects.filter(user=superuser, action='xlsx_export').exists()
        )

    def test_xlsx_escapes_formula_like_employee_name(self):
        self.employee.first_name = '=HYPERLINK("https://invalid.example")'
        self.employee.last_name = ''
        self.employee.save(update_fields=['first_name', 'last_name'])
        self.client.login(username=self.officer.username, password=self.password)

        response = self.client.get(reverse('reports:training_registry_xlsx'))

        self.assertEqual(response.status_code, 200)
        workbook = load_workbook(BytesIO(response.content), data_only=False)
        self.assertEqual(
            workbook['Реестр обучения']['B6'].value,
            "'=HYPERLINK(\"https://invalid.example\")",
        )

    def test_control_characters_do_not_break_xlsx_or_docx(self):
        self.course.title = 'Основы\x01 информационной безопасности'
        self.course.save(update_fields=['title'])
        self.client.login(username=self.officer.username, password=self.password)

        xlsx_response = self.client.get(reverse('reports:training_registry_xlsx'))
        docx_response = self.client.get(
            reverse('reports:assignment_certificate_docx', args=[self.assignment.pk])
        )

        self.assertEqual(xlsx_response.status_code, 200)
        self.assertEqual(docx_response.status_code, 200)

        workbook = load_workbook(BytesIO(xlsx_response.content), data_only=False)
        self.assertEqual(
            workbook['Реестр обучения']['F6'].value,
            'Основы информационной безопасности',
        )
        document = Document(BytesIO(docx_response.content))
        document_text = '\n'.join(p.text for p in document.paragraphs)
        self.assertIn('Основы информационной безопасности', document_text)
        self.assertNotIn('\x01', document_text)
